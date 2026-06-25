from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent
OUT = ROOT / "Rapport_CICD_Django.docx"

BLUE = "0D6EFD"
NAVY = "163A63"
PALE_BLUE = "EAF2FF"
GREEN = "DFF4E5"
GREEN_TEXT = "176B35"
RED = "FDE7E7"
RED_TEXT = "A61B1B"
GRAY = "667085"
LIGHT = "F5F7FA"


def set_cell_fill(cell, color):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    cell._tc.get_or_add_tcPr().append(shd)


def set_cell_margins(cell, top=100, start=140, bottom=100, end=140):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_font(run, size=10, bold=False, color="1F2937", name="Calibri"):
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def add_text(doc, text, size=10, color="1F2937", bold=False, after=5, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.15
    if align is not None:
        p.alignment = align
    set_font(p.add_run(text), size=size, bold=bold, color=color)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    set_font(p.add_run(text), size=15 if level == 1 else 12, bold=True, color=NAVY)
    return p


def add_status(doc, label, text, good=True):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = (Inches(1.3), Inches(5.0))
    for cell, width in zip(table.rows[0].cells, widths):
        cell.width = width
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(cell)
        set_cell_fill(cell, GREEN if good else RED)
    p0 = table.cell(0, 0).paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p0.add_run(label), size=9, bold=True, color=GREEN_TEXT if good else RED_TEXT)
    set_font(table.cell(0, 1).paragraphs[0].add_run(text), size=9.5, color="1F2937")
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_capture_placeholder(doc, title, instruction, height_lines=3):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Inches(6.3)
    set_cell_fill(cell, LIGHT)
    set_cell_margins(cell, top=180, start=180, bottom=180, end=180)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run(title + "\n"), size=9.5, bold=True, color=NAVY)
    set_font(p.add_run(instruction + ("\n" * height_lines)), size=8.5, color=GRAY)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.65)
section.bottom_margin = Inches(0.65)
section.left_margin = Inches(0.85)
section.right_margin = Inches(0.85)
section.header_distance = Inches(0.3)
section.footer_distance = Inches(0.3)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10)
normal.paragraph_format.space_after = Pt(5)
normal.paragraph_format.line_spacing = 1.15
for level, size, before, after in ((1, 15, 10, 5), (2, 12, 7, 4)):
    style = styles[f"Heading {level}"]
    style.font.name = "Calibri"
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor.from_string(NAVY)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)

header = section.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_font(header.add_run("RATT DEVOPS  |  RAPPORT CI/CD"), size=8, bold=True, color=GRAY)
footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(footer.add_run("Django • GitHub Actions • Vercel"), size=8, color=GRAY)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(2)
set_font(p.add_run("RAPPORT TECHNIQUE"), size=9, bold=True, color=BLUE)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(4)
set_font(p.add_run("Pipeline CI/CD Django"), size=25, bold=True, color=NAVY)
add_text(doc, "Contrôle automatique de la couleur de fond et déploiement Vercel", size=12, color=GRAY, after=12)

add_status(doc, "ÉTAT LOCAL", "Tests validés avec fond bleu ; échec reproduit avec fond orange.", good=True)

add_heading(doc, "1. Architecture mise en place", 1)
add_text(doc, "L'application Django contient une vue unique et un template HTML minimaliste. La couleur de fond est définie sur body. Le script tests/test_background_color.py lit cette déclaration et refuse explicitement orange, #ffa500, #ff8c00 et leurs équivalents RGB.")

table = doc.add_table(rows=4, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = False
rows = [
    ("Branche test", "Branche de travail et de push par défaut"),
    ("Pull request", "test → main ; exécute le contrôle avant fusion"),
    ("Branche main", "Production ; protégée par le statut Tests obligatoires"),
    ("Vercel", "Déploiement automatique uniquement depuis main"),
]
for i, (label, value) in enumerate(rows):
    for j, text in enumerate((label, value)):
        cell = table.cell(i, j)
        cell.width = Inches(1.55 if j == 0 else 4.75)
        set_cell_margins(cell)
        set_cell_fill(cell, PALE_BLUE if j == 0 else "FFFFFF")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_font(cell.paragraphs[0].add_run(text), size=9, bold=(j == 0), color=NAVY if j == 0 else "1F2937")

add_heading(doc, "2. Paramètres GitHub indispensables", 1)
add_text(doc, "Dans Settings → Branches, protéger main avec « Require a pull request before merging » et « Require status checks to pass before merging », puis sélectionner le statut Tests obligatoires. Sans cette règle, GitHub Actions signale l'échec mais ne bloque pas strictement la fusion.")
add_text(doc, "Le workflow est déclenché sur pull_request vers main et sur push vers main. Le premier événement bloque la fusion ; le second vérifie aussi l'état réel de la production.", size=9.5, color=GRAY)

doc.add_page_break()

add_text(doc, "SCÉNARIO 1", size=9, bold=True, color=GREEN_TEXT, after=2)
add_heading(doc, "Succès — fond bleu", 1)
add_text(doc, "1. Le template conserve background-color: #0d6efd. Le développeur pousse le commit sur test, puis ouvre une pull request vers main.")
add_capture_placeholder(doc, "CAPTURE A — Push sur test + page bleue", "Insérer le terminal montrant « git push origin test » et la page locale/preview bleue.", 2)

add_text(doc, "2. GitHub Actions installe Python 3.12, exécute le test de couleur puis les tests Django. Les deux commandes terminent avec OK.")
add_status(doc, "SUCCÈS", "test_background_is_not_orange ... ok  |  test_home_page_is_available ... ok", good=True)
add_capture_placeholder(doc, "CAPTURE B — GitHub Actions vert", "Insérer l'onglet Checks/Actions avec le job Tests obligatoires réussi.", 2)

add_text(doc, "3. La pull request devient fusionnable. Après merge, Vercel détecte main et publie l'application.")
add_capture_placeholder(doc, "CAPTURE C — Merge + Vercel", "Insérer la PR fusionnée, puis le site Vercel affichant le fond bleu et son URL.", 3)

add_text(doc, "Résultat attendu : main contient la version bleue ; le déploiement Vercel est fonctionnel.", size=10, bold=True, color=GREEN_TEXT, after=0)

doc.add_page_break()

add_text(doc, "SCÉNARIO 2", size=9, bold=True, color=RED_TEXT, after=2)
add_heading(doc, "Échec provoqué — fond orange", 1)
add_text(doc, "1. Sur test, remplacer temporairement #0d6efd par orange, committer et pousser. La pull request vers main relance automatiquement le contrôle.")
add_capture_placeholder(doc, "CAPTURE D — Modification orange", "Insérer le diff du template et la page orange sur la branche test.", 2)

add_text(doc, "2. Le test détecte la valeur interdite et termine avec un code de sortie 1. L'échec a été reproduit localement avant remise du rapport :")
add_status(doc, "ÉCHEC", "AssertionError: Déploiement bloqué : le fond 'orange' est orange.", good=False)
add_capture_placeholder(doc, "CAPTURE E — GitHub Actions rouge", "Insérer le log du job Vérifier la couleur de fond avec FAILED (failures=1).", 2)

add_text(doc, "3. La protection de main affiche un contrôle requis en échec : le bouton de fusion est bloqué. Puisque main n'est pas modifiée, aucun déploiement de production Vercel ne doit partir.")
add_capture_placeholder(doc, "CAPTURE F — Blocage strict", "Insérer la PR avec Required check failed / Merge blocked et l'absence de nouveau déploiement Production sur Vercel.", 3)

add_status(doc, "CONCLUSION", "Le mécanisme protège la production avant la fusion, à condition que la règle de protection de main soit activée.", good=True)
add_text(doc, "Note de rendu : les cadres A–F doivent être remplacés par les captures du dépôt GitHub public et du projet Vercel de l'étudiant. Ces preuves dépendent des comptes externes et ne peuvent pas être fabriquées localement.", size=8.5, color=GRAY, after=0)

doc.save(OUT)
print(OUT)
